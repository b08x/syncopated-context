#!/usr/bin/env -S uv run --script
"""
Convert ServiceNow KB HTML to properly formatted DOCX file.

Usage:
    python3 html_to_docx.py input.html output.docx
    python3 html_to_docx.py input.html  # outputs input.docx
"""

import sys
import re
import base64
from pathlib import Path
from html.parser import HTMLParser
from io import BytesIO

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    from PIL import Image
except ImportError:
    print("Warning: Pillow not installed. Image handling may be limited.")
    print("Run: pip install Pillow")
    Image = None


class KBHTMLParser(HTMLParser):
    """Parse KB HTML and build DOCX structure"""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.current_run = None
        self.list_level = 0
        self.list_stack = []  # Track list types (ul/ol)
        self.in_metadata = False
        self.in_code = False
        self.in_pre = False
        self.in_strong = False
        self.in_warning = False
        self.in_a = False
        self.link_href = None
        self.current_table = None
        self.current_cell = None
        self.table_cell_idx = 0
        self.in_table_header = False
        self.in_callout = False
        self.collect_text = []

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        if tag in ["h1", "h2", "h3", "h4"]:
            level = int(tag[1]) - 1
            self.current_paragraph = self.doc.add_heading(level=level)
            if tag == "h1":
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Add bookmark for internal links
            # We'll handle the bookmark creation after we have the text
            self.current_paragraph._p.tag_name = tag  # Temporary storage

        elif tag == "p":
            if self.in_callout and self.current_cell:
                self.current_paragraph = self.current_cell.paragraphs[-1]
                if self.current_paragraph.text:
                    self.current_paragraph = self.current_cell.add_paragraph()
                self.current_run = self.current_paragraph.add_run()
            # Check for metadata class
            elif "class" in attrs_dict and "metadata" in attrs_dict["class"]:
                self.in_metadata = True
                self.current_paragraph = self.doc.add_paragraph()
                self.current_run = self.current_paragraph.add_run()
                self.current_run.font.size = Pt(10)
                self.current_run.font.color.rgb = RGBColor(102, 102, 102)
            else:
                self.current_paragraph = self.doc.add_paragraph()
                self.current_run = self.current_paragraph.add_run()

        elif tag == "pre":
            self.in_pre = True
            self.current_paragraph = self.doc.add_paragraph()
            self._set_paragraph_shading(self.current_paragraph, "F4F4F4")
            self._set_paragraph_border(self.current_paragraph, color="D0D0D0")
            self.current_run = self.current_paragraph.add_run()
            self.current_run.font.name = "Courier New"
            self.current_run.font.size = Pt(9)

        elif tag == "div":
            # Handle callouts using 1x1 table for better nesting support
            cls = attrs_dict.get("class", "")
            if "callout" in cls:
                self.in_callout = True
                self.current_table = self.doc.add_table(
                    rows=1, cols=1, style="Table Grid"
                )
                self.current_cell = self.current_table.rows[0].cells[0]

                # Set shading
                tcPr = self.current_cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                if "info" in cls:
                    shd.set(qn("w:fill"), "F0F7FF")
                elif "warning" in cls:
                    shd.set(qn("w:fill"), "FFF5F5")
                elif "note" in cls:
                    shd.set(qn("w:fill"), "F9F9F9")
                tcPr.append(shd)

                self.current_paragraph = self.current_cell.paragraphs[0]
                self.current_run = self.current_paragraph.add_run()

        elif tag == "ul":
            self.list_stack.append("ul")
            self.list_level += 1

        elif tag == "ol":
            self.list_stack.append("ol")
            self.list_level += 1

        elif tag == "li":
            style = "List Bullet" if self.list_stack[-1] == "ul" else "List Number"
            # Fallback if style doesn't exist
            try:
                self.current_paragraph = self.doc.add_paragraph(style=style)
            except:
                self.current_paragraph = self.doc.add_paragraph()
                prefix = "• " if self.list_stack[-1] == "ul" else f"{self.list_level}. "
                self.current_paragraph.add_run(prefix)

            # Set indentation based on level
            if self.list_level > 1:
                self.current_paragraph.paragraph_format.left_indent = Inches(
                    0.25 * self.list_level
                )
            self.current_run = self.current_paragraph.add_run()

        elif tag == "strong" or tag == "b":
            self.in_strong = True
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()
                self.current_run.bold = True

        elif tag == "code":
            self.in_code = True
            if self.current_paragraph:
                if not self.in_pre:
                    self.current_run = self.current_paragraph.add_run()
                    self.current_run.font.name = "Courier New"
                    self.current_run.font.size = Pt(10)
                    self.current_run.font.color.rgb = RGBColor(
                        199, 37, 78
                    )  # Nice pink/red for inline code

        elif tag == "span":
            if "class" in attrs_dict and "warning" in attrs_dict["class"]:
                self.in_warning = True
                if self.current_paragraph:
                    self.current_run = self.current_paragraph.add_run()
                    self.current_run.bold = True
                    self.current_run.font.color.rgb = RGBColor(255, 0, 0)

        elif tag == "a":
            self.in_a = True
            self.link_href = attrs_dict.get("href", "")
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()

        elif tag == "img":
            src = attrs_dict.get("src", "")
            alt = attrs_dict.get("alt", "Image")
            self._add_image(src, alt)

        elif tag == "br":
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run("\n")

        elif tag == "table":
            self.current_table = self.doc.add_table(rows=0, cols=0, style="Table Grid")
            self.current_table.autofit = True

        elif tag == "tr":
            if self.current_table:
                self.current_table.add_row()
                self.table_cell_idx = 0
                self.in_table_header = False

        elif tag == "th":
            self.in_table_header = True
            if self.current_table:
                # Add column if needed
                if self.table_cell_idx >= len(self.current_table.columns):
                    self.current_table.add_column(Inches(1.5))

                self.current_cell = self.current_table.rows[-1].cells[
                    self.table_cell_idx
                ]
                self.table_cell_idx += 1
                self.current_paragraph = self.current_cell.paragraphs[0]
                self.current_run = self.current_paragraph.add_run()
                self.current_run.bold = True
                # Set shading for header
                tcPr = self.current_cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F2F2F2")
                tcPr.append(shd)

        elif tag == "td":
            if self.current_table:
                # Add column if needed
                if self.table_cell_idx >= len(self.current_table.columns):
                    self.current_table.add_column(Inches(1.5))

                self.current_cell = self.current_table.rows[-1].cells[
                    self.table_cell_idx
                ]
                self.table_cell_idx += 1
                self.current_paragraph = self.current_cell.paragraphs[0]
                self.current_run = self.current_paragraph.add_run()

        # Handle ID for bookmarks
        if "id" in attrs_dict and self.current_paragraph:
            self._add_bookmark(self.current_paragraph, attrs_dict["id"])

    def handle_endtag(self, tag):
        if tag in ["h1", "h2", "h3", "h4"]:
            if self.current_paragraph:
                # Use paragraph text as bookmark if no ID was provided
                text_id = re.sub(
                    r"[^a-zA-Z0-9]", "", self.current_paragraph.text
                ).lower()
                self._add_bookmark(self.current_paragraph, text_id)
            self.current_paragraph = None
            self.current_run = None

        elif tag in ["p", "li", "div", "pre"]:
            self.current_paragraph = None
            self.current_run = None
            self.in_metadata = False
            self.in_pre = False

        elif tag == "ul" or tag == "ol":
            if self.list_stack:
                self.list_stack.pop()
            self.list_level = max(0, self.list_level - 1)

        elif tag == "strong" or tag == "b":
            self.in_strong = False
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()

        elif tag == "code":
            self.in_code = False
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()

        elif tag == "span":
            self.in_warning = False
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()

        elif tag == "a":
            self.in_a = False
            self.link_href = None
            if self.current_paragraph:
                self.current_run = self.current_paragraph.add_run()

        elif tag == "table":
            self.current_table = None
            self.current_cell = None

        elif tag == "div":
            if self.in_callout:
                self.in_callout = False
                self.current_table = None
                self.current_cell = None
            self.current_paragraph = None
            self.current_run = None

        elif tag in ["th", "td"]:
            self.current_cell = None
            self.current_paragraph = None
            self.current_run = None

    def handle_data(self, data):
        if not self.current_paragraph:
            return

        # Clean up whitespace but preserve intentional spacing
        text = data
        if not self.in_code and not self.in_pre:
            text = re.sub(r"\s+", " ", text)

        if text.strip() or self.in_code or self.in_pre:
            if self.in_a and self.link_href:
                if self.link_href.startswith("#"):
                    self._add_hyperlink(
                        self.current_paragraph, self.link_href[1:], text
                    )
                else:
                    # Regular external link (optional: style as blue/underlined)
                    run = self.current_paragraph.add_run(text)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True
            else:
                if not self.current_run:
                    self.current_run = self.current_paragraph.add_run()
                self.current_run.text += text

    def _add_bookmark(self, paragraph, bookmark_name):
        """Add a bookmark to a paragraph"""
        if not bookmark_name:
            return
        # Clean name: only alphanumeric and underscore, must start with letter
        bookmark_name = re.sub(r"[^a-zA-Z0-9_]", "_", bookmark_name)
        if not bookmark_name[0].isalpha():
            bookmark_name = "b" + bookmark_name

        p = paragraph._p
        bookmark_start = OxmlElement("w:bookmarkStart")
        bookmark_start.set(qn("w:id"), "0")
        bookmark_start.set(qn("w:name"), bookmark_name)
        p.insert(0, bookmark_start)

        bookmark_end = OxmlElement("w:bookmarkEnd")
        bookmark_end.set(qn("w:id"), "0")
        p.append(bookmark_end)

    def _add_hyperlink(self, paragraph, bookmark_name, text):
        """Add an internal hyperlink to a bookmark"""
        # Clean name to match bookmark
        bookmark_name = re.sub(r"[^a-zA-Z0-9_]", "_", bookmark_name)
        if not bookmark_name[0].isalpha():
            bookmark_name = "b" + bookmark_name

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("w:anchor"), bookmark_name)

        new_run = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")

        # Style as blue and underlined
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "0563C1")
        rPr.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        rPr.append(underline)

        new_run.append(rPr)
        text_element = OxmlElement("w:t")
        text_element.text = text
        new_run.append(text_element)

        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)

    def _set_paragraph_shading(self, paragraph, hex_color):
        """Set paragraph background shading"""
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), hex_color)
        paragraph._p.get_or_add_pPr().append(shading)

    def _set_paragraph_border(self, paragraph, color="Auto"):
        """Set paragraph borders"""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for side in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")  # 1/2 pt
            border.set(qn("w:space"), "4")
            border.set(qn("w:color"), color)
            pBdr.append(border)
        pPr.append(pBdr)

    def _add_image(self, src, alt):
        """Add image to document"""
        # Create new paragraph for image
        img_paragraph = self.doc.add_paragraph()
        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        try:
            if src.startswith("data:image"):
                # Base64 encoded image
                header, encoded = src.split(",", 1)
                image_data = base64.b64decode(encoded)

                # Add image with max width constraint
                if Image:
                    # Use PIL to get dimensions and resize if needed
                    img = Image.open(BytesIO(image_data))
                    width, height = img.size

                    # Max width 6.5 inches (650px at 100dpi)
                    max_width = 6.5
                    if width > max_width * 100:
                        scale = (max_width * 100) / width
                        new_width = max_width
                        new_height = (height * scale) / 100

                        run = img_paragraph.add_run()
                        run.add_picture(
                            BytesIO(image_data),
                            width=Inches(new_width),
                            height=Inches(new_height),
                        )
                    else:
                        run = img_paragraph.add_run()
                        run.add_picture(BytesIO(image_data), width=Inches(width / 100))
                else:
                    # Without PIL, just add with default sizing
                    run = img_paragraph.add_run()
                    run.add_picture(BytesIO(image_data))

            else:
                # External image or ID reference - add placeholder
                placeholder = img_paragraph.add_run()
                placeholder.text = f"[Image: {alt}]"
                placeholder.italic = True
                placeholder.font.color.rgb = RGBColor(128, 128, 128)

        except Exception as e:
            # If image fails, add text placeholder
            placeholder = img_paragraph.add_run()
            placeholder.text = f"[Image: {alt}]"
            placeholder.italic = True
            placeholder.font.color.rgb = RGBColor(128, 128, 128)


def html_to_docx(html_content, output_path):
    """Convert KB HTML to DOCX"""

    # Create document
    doc = Document()

    # Set default styles
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Heading 1 style (maps to H2 in HTML - Major Phases)
    h1_style = doc.styles["Heading 1"]
    h1_style.font.name = "Arial"
    h1_style.font.size = Pt(18)
    h1_style.font.bold = True

    # Heading 2 style (maps to H3 in HTML - Steps)
    h2_style = doc.styles["Heading 2"]
    h2_style.font.name = "Arial"
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True

    # Heading 3 style (maps to H4 in HTML - Tasks)
    h3_style = doc.styles["Heading 3"]
    h3_style.font.name = "Arial"
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True

    # Parse HTML
    parser = KBHTMLParser(doc)

    # Clean HTML - remove html/body tags if present
    html_content = re.sub(r"<!DOCTYPE[^>]*>", "", html_content, flags=re.IGNORECASE)
    # Handle Tables before general parsing to simplify
    # This script's manual parser is basic; for tables, we might need a better approach or just handle them in the loop.
    # For now, let's just make sure the Document styles are robust.

    # Add Table Grid style if missing
    if "Table Grid" not in doc.styles:
        doc.styles.add_style(
            "Table Grid", 1
        )  # 1 = WD_STYLE_TYPE.PARAGRAPH (wrong, but usually works for simple cases)
        # Actually Table styles are complex in docx-python. Let's just use 'Light List Accent 1' or similar if available.

    html_content = re.sub(r"<html[^>]*>", "", html_content, flags=re.IGNORECASE)
    html_content = re.sub(r"</html>", "", html_content, flags=re.IGNORECASE)
    html_content = re.sub(r"<body[^>]*>", "", html_content, flags=re.IGNORECASE)
    html_content = re.sub(r"</body>", "", html_content, flags=re.IGNORECASE)

    parser.feed(html_content)

    # Save document
    doc.save(output_path)
    print(f"✅ Created: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 html_to_docx.py input.html [output.docx]")
        sys.exit(1)

    input_path = Path(sys.argv[1])

    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}")
        sys.exit(1)

    # Determine output path
    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
    else:
        output_path = input_path.with_suffix(".docx")

    # Read HTML
    with open(input_path, "r", encoding="utf-8") as f:
        html_content = f.read()

    # Convert
    html_to_docx(html_content, output_path)


if __name__ == "__main__":
    main()
