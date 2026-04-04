#!/usr/bin/env python3
"""
Resume DOCX Generator

Generates Word documents from structured resume data with template-aware styling.
Supports all four resume templates: minimal, professional, modern, executive.

Usage:
    python3 generate_docx.py --template minimal --data resume_data.json output.docx
    
    Or use programmatically:
        from generate_docx import ResumeGenerator
        generator = ResumeGenerator(template='minimal')
        generator.generate(resume_data, 'output.docx')
"""

import sys
import json
import argparse
from pathlib import Path
from dataclasses import dataclass
from typing import Dict, List, Optional, Any
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx not installed. Install with:")
    print("  pip install python-docx --break-system-packages")
    sys.exit(1)


@dataclass
class StyleConfig:
    """Style configuration matching HTML templates"""
    accent_color: str = "#2563eb"  # Hex color
    font_family: str = "system"    # system | serif | sans | mono
    font_size: str = "medium"      # small | medium | large
    spacing: str = "normal"        # compact | normal | relaxed
    page_size: str = "letter"      # letter | a4
    
    def get_rgb_color(self) -> RGBColor:
        """Convert hex color to RGBColor"""
        hex_color = self.accent_color.lstrip('#')
        return RGBColor(*[int(hex_color[i:i+2], 16) for i in (0, 2, 4)])
    
    def get_font_name(self) -> str:
        """Map font family to actual font name"""
        mapping = {
            'system': 'Calibri',
            'serif': 'Georgia',
            'sans': 'Arial',
            'mono': 'Consolas'
        }
        return mapping.get(self.font_family, 'Calibri')
    
    def get_base_font_size(self) -> int:
        """Get base font size in points"""
        mapping = {
            'small': 9,
            'medium': 10,
            'large': 11
        }
        return mapping.get(self.font_size, 10)
    
    def get_spacing_values(self) -> Dict[str, float]:
        """Get spacing values in inches"""
        mapping = {
            'compact': {'section': 0.15, 'item': 0.08, 'after': 6},
            'normal': {'section': 0.25, 'item': 0.12, 'after': 8},
            'relaxed': {'section': 0.35, 'item': 0.18, 'after': 10}
        }
        return mapping.get(self.spacing, mapping['normal'])


class ResumeGenerator:
    """Generate Word documents from resume data"""
    
    def __init__(self, template: str = 'minimal', style: Optional[StyleConfig] = None):
        self.template = template
        self.style = style or StyleConfig()
        self.doc = Document()
        self._setup_document()
        self._setup_styles()
    
    def _setup_document(self):
        """Configure document-level settings"""
        section = self.doc.sections[0]
        
        # Set margins
        if self.template == 'executive':
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)
        else:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
    
    def _setup_styles(self):
        """Create custom styles matching HTML templates"""
        styles = self.doc.styles
        
        # Name style (h1 equivalent)
        name_style = styles.add_style('ResumeName', WD_STYLE_TYPE.PARAGRAPH)
        name_font = name_style.font
        name_font.name = self.style.get_font_name()
        name_font.size = Pt(self.style.get_base_font_size() * 1.8)
        name_font.bold = True
        name_font.color.rgb = RGBColor(17, 17, 17) if self.template != 'professional' else self.style.get_rgb_color()
        name_style.paragraph_format.space_after = Pt(4)
        
        if self.template in ['minimal', 'executive']:
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contact info style
        contact_style = styles.add_style('ResumeContact', WD_STYLE_TYPE.PARAGRAPH)
        contact_font = contact_style.font
        contact_font.name = self.style.get_font_name()
        contact_font.size = Pt(self.style.get_base_font_size() * 0.9)
        contact_font.color.rgb = RGBColor(85, 85, 85)
        contact_style.paragraph_format.space_after = Pt(self.style.get_spacing_values()['after'])
        
        if self.template in ['minimal', 'executive']:
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Section heading style (h2 equivalent)
        heading_style = styles.add_style('ResumeSection', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = self.style.get_font_name()
        heading_font.size = Pt(self.style.get_base_font_size() * 1.1)
        heading_font.bold = True
        heading_font.color.rgb = self.style.get_rgb_color()
        heading_style.paragraph_format.space_before = Pt(self.style.get_spacing_values()['after'] * 1.5)
        heading_style.paragraph_format.space_after = Pt(6)
        heading_style.paragraph_format.keep_with_next = True
        
        if self.template == 'executive':
            heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_font.size = Pt(self.style.get_base_font_size())
            heading_font.bold = False
            heading_font.all_caps = True
        
        # Job title / entry header style
        title_style = styles.add_style('ResumeTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = self.style.get_font_name()
        title_font.size = Pt(self.style.get_base_font_size())
        title_font.bold = True
        title_font.color.rgb = RGBColor(34, 34, 34)
        title_style.paragraph_format.space_after = Pt(2)
        title_style.paragraph_format.keep_with_next = True
        
        # Date/meta style
        meta_style = styles.add_style('ResumeMeta', WD_STYLE_TYPE.CHARACTER)
        meta_font = meta_style.font
        meta_font.name = self.style.get_font_name()
        meta_font.size = Pt(self.style.get_base_font_size() * 0.9)
        meta_font.color.rgb = RGBColor(102, 102, 102)
        
        # Body/bullet style
        body_style = styles.add_style('ResumeBody', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = self.style.get_font_name()
        body_font.size = Pt(self.style.get_base_font_size())
        body_style.paragraph_format.space_after = Pt(3)
        body_style.paragraph_format.left_indent = Inches(0.25)
    
    def _add_header(self, data: Dict[str, Any]):
        """Add resume header with name and contact info"""
        # Name
        name_para = self.doc.add_paragraph(data.get('name', ''), style='ResumeName')
        
        # Contact info
        contact_items = []
        if 'email' in data:
            contact_items.append(data['email'])
        if 'phone' in data:
            contact_items.append(data['phone'])
        if 'location' in data:
            contact_items.append(data['location'])
        if 'linkedin' in data:
            contact_items.append(data['linkedin'])
        
        if contact_items:
            separator = ' · ' if self.template == 'minimal' else ' | '
            contact_para = self.doc.add_paragraph(separator.join(contact_items), style='ResumeContact')
        
        # Add border for professional template
        if self.template == 'professional':
            name_para.paragraph_format.border_bottom_color = self.style.get_rgb_color()
            name_para.paragraph_format.border_bottom_width = Pt(2)
    
    def _add_summary(self, summary: str):
        """Add professional summary"""
        if not summary:
            return
        
        para = self.doc.add_paragraph()
        run = para.add_run(summary)
        run.font.name = self.style.get_font_name()
        run.font.size = Pt(self.style.get_base_font_size())
        run.font.italic = True
        run.font.color.rgb = RGBColor(68, 68, 68)
        
        if self.template in ['minimal', 'executive']:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        para.paragraph_format.space_after = Pt(self.style.get_spacing_values()['after'])
    
    def _add_section_heading(self, title: str):
        """Add section heading"""
        para = self.doc.add_paragraph(title.upper(), style='ResumeSection')
        
        # Add underline for professional template
        if self.template == 'professional':
            para.paragraph_format.border_bottom_color = RGBColor(208, 208, 208)
            para.paragraph_format.border_bottom_width = Pt(1)
    
    def _add_experience_entry(self, entry: Dict[str, Any]):
        """Add work experience entry"""
        # Title and company
        title_para = self.doc.add_paragraph(style='ResumeTitle')
        title_para.add_run(entry.get('title', ''))
        
        if 'company' in entry:
            title_para.add_run(' — ')
            org_run = title_para.add_run(entry['company'])
            org_run.font.bold = False
            org_run.font.color.rgb = RGBColor(68, 68, 68)
        
        # Date range
        if 'date_range' in entry:
            title_para.add_run('\t')
            date_run = title_para.add_run(entry['date_range'])
            date_run.font.bold = False
            date_run.font.size = Pt(self.style.get_base_font_size() * 0.9)
            date_run.font.color.rgb = RGBColor(102, 102, 102)
        
        # Bullets
        for bullet in entry.get('bullets', []):
            bullet_para = self.doc.add_paragraph(bullet, style='ResumeBody')
            bullet_para.paragraph_format.left_indent = Inches(0.25)
            bullet_para.style = 'List Bullet'
        
        # Add spacing after entry
        self.doc.add_paragraph().paragraph_format.space_after = Pt(self.style.get_spacing_values()['item'] * 10)
    
    def _add_education_entry(self, entry: Dict[str, Any]):
        """Add education entry"""
        title_para = self.doc.add_paragraph(style='ResumeTitle')
        title_para.add_run(entry.get('degree', ''))
        
        if 'institution' in entry:
            title_para.add_run(' — ')
            inst_run = title_para.add_run(entry['institution'])
            inst_run.font.bold = False
            inst_run.font.color.rgb = RGBColor(68, 68, 68)
        
        if 'year' in entry:
            title_para.add_run('\t')
            year_run = title_para.add_run(str(entry['year']))
            year_run.font.bold = False
            year_run.font.size = Pt(self.style.get_base_font_size() * 0.9)
            year_run.font.color.rgb = RGBColor(102, 102, 102)
        
        if 'details' in entry:
            detail_para = self.doc.add_paragraph(entry['details'], style='ResumeBody')
        
        self.doc.add_paragraph().paragraph_format.space_after = Pt(self.style.get_spacing_values()['item'] * 10)
    
    def _add_skills_section(self, skills: List[Any]):
        """Add skills section"""
        if isinstance(skills, dict):
            # Categorized skills
            for category, skill_list in skills.items():
                para = self.doc.add_paragraph(style='ResumeBody')
                para.paragraph_format.left_indent = Inches(0)
                
                cat_run = para.add_run(f"{category}: ")
                cat_run.font.bold = True
                cat_run.font.color.rgb = self.style.get_rgb_color()
                
                skills_text = ', '.join(skill_list) if isinstance(skill_list, list) else str(skill_list)
                para.add_run(skills_text)
        else:
            # Simple list
            skills_text = ', '.join(skills) if isinstance(skills, list) else str(skills)
            para = self.doc.add_paragraph(skills_text, style='ResumeBody')
            para.paragraph_format.left_indent = Inches(0)
            
            if self.template == 'executive':
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def generate(self, data: Dict[str, Any], output_path: Path) -> bool:
        """
        Generate DOCX from resume data.
        
        Args:
            data: Resume data dictionary with sections
            output_path: Path to output DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Header
            self._add_header(data)
            
            # Summary
            if 'summary' in data:
                self._add_summary(data['summary'])
            
            # Experience
            if 'experience' in data:
                self._add_section_heading('Experience')
                for entry in data['experience']:
                    self._add_experience_entry(entry)
            
            # Education
            if 'education' in data:
                self._add_section_heading('Education')
                for entry in data['education']:
                    self._add_education_entry(entry)
            
            # Skills
            if 'skills' in data:
                self._add_section_heading('Skills')
                self._add_skills_section(data['skills'])
            
            # Projects
            if 'projects' in data:
                self._add_section_heading('Projects')
                for entry in data['projects']:
                    self._add_experience_entry(entry)  # Same format as experience
            
            # Certifications
            if 'certifications' in data:
                self._add_section_heading('Certifications')
                for cert in data['certifications']:
                    para = self.doc.add_paragraph(style='ResumeBody')
                    para.paragraph_format.left_indent = Inches(0.25)
                    para.style = 'List Bullet'
                    
                    name_run = para.add_run(cert.get('name', ''))
                    name_run.font.bold = True
                    
                    if 'issuer' in cert or 'date' in cert:
                        para.add_run(' — ')
                        details = []
                        if 'issuer' in cert:
                            details.append(cert['issuer'])
                        if 'date' in cert:
                            details.append(cert['date'])
                        para.add_run(', '.join(details))
            
            # Save document
            print(f"Generating DOCX: {output_path}")
            self.doc.save(str(output_path))
            print(f"✓ DOCX generated successfully: {output_path}")
            print(f"  Size: {output_path.stat().st_size / 1024:.1f} KB")
            return True
            
        except Exception as e:
            print(f"✗ Error generating DOCX: {e}", file=sys.stderr)
            import traceback
            traceback.print_exc()
            return False


def main():
    parser = argparse.ArgumentParser(
        description='Generate DOCX from resume data with template styling'
    )
    parser.add_argument('--data', type=Path, required=True, help='Resume data JSON file')
    parser.add_argument('output', type=Path, help='Output DOCX file')
    parser.add_argument(
        '--template',
        choices=['minimal', 'professional', 'modern', 'executive'],
        default='minimal',
        help='Resume template to use'
    )
    parser.add_argument('--accent-color', help='Accent color (hex, e.g., #2563eb)')
    parser.add_argument(
        '--font-family',
        choices=['system', 'serif', 'sans', 'mono'],
        help='Font family'
    )
    parser.add_argument(
        '--font-size',
        choices=['small', 'medium', 'large'],
        help='Font size'
    )
    parser.add_argument(
        '--spacing',
        choices=['compact', 'normal', 'relaxed'],
        help='Spacing'
    )
    
    args = parser.parse_args()
    
    # Validate data file
    if not args.data.exists():
        print(f"Error: Data file not found: {args.data}", file=sys.stderr)
        sys.exit(1)
    
    # Load resume data
    try:
        with open(args.data) as f:
            data = json.load(f)
    except Exception as e:
        print(f"Error loading data file: {e}", file=sys.stderr)
        sys.exit(1)
    
    # Build style config
    style = StyleConfig()
    if args.accent_color:
        style.accent_color = args.accent_color
    if args.font_family:
        style.font_family = args.font_family
    if args.font_size:
        style.font_size = args.font_size
    if args.spacing:
        style.spacing = args.spacing
    
    # Ensure output directory exists
    args.output.parent.mkdir(parents=True, exist_ok=True)
    
    # Generate DOCX
    generator = ResumeGenerator(template=args.template, style=style)
    success = generator.generate(data, args.output)
    
    sys.exit(0 if success else 1)


if __name__ == '__main__':
    main()
